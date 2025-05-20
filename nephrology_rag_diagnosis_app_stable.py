
import streamlit as st
from io import BytesIO
from docx import Document
from sentence_transformers import SentenceTransformer
import faiss
import json

# 유효성 검사 함수
def validate_input(eGFR, creatinine, albumin, proteinuria):
    errors = []
    if not (0 <= eGFR <= 150):
        errors.append("eGFR 값은 0~150 사이여야 합니다.")
    if not (0 <= creatinine <= 20):
        errors.append("Creatinine 값은 0~20 사이여야 합니다.")
    if not (0 <= albumin <= 6):
        errors.append("Albumin 값은 0~6 사이여야 합니다.")
    if not (0 <= proteinuria <= 5000):
        errors.append("단백뇨 값은 0~5000 사이여야 합니다.")
    return errors

# Streamlit 설정
st.set_page_config(page_title="신장내과 진단 시스템", layout="centered")
st.title("🧜 신장내과 질병 예측 시스템")

st.markdown("""
이 시스템은 신장내과 질병군(예: CKD, AKI, Nephrotic Syndrome 등)에 대해
혈액검사 수치를 기반으로 질병을 보조하고,
질병에 관한 질문에 규격 기반 응답을 제공하는 **AI 기반 질병 지원 도구**입니다.
""")

st.header("🦢 검사 수치 입력")

with st.form("manual_input_form"):
    eGFR = st.number_input("eGFR (mL/min/1.73m²)", min_value=0.0, max_value=200.0, step=0.1)
    creatinine = st.number_input("Creatinine (mg/dL)", min_value=0.0, max_value=20.0, step=0.1)
    albumin = st.number_input("Albumin (g/dL)", min_value=0.0, max_value=6.0, step=0.1)
    proteinuria = st.number_input("단백뇨 (mg/dL)", min_value=0.0, max_value=5000.0, step=1.0)
    submitted = st.form_submit_button("질병 결과 확인")

def generate_report(user_question, matched_docs, llm_answer, report_type="질병예측"):
    doc = Document()
    doc.add_heading("RAG 기반 질의응답 결과 리포트", level=1)
    doc.add_paragraph(f"자동 생성 질의: {user_question}")
    doc.add_heading("검사된 문서", level=2)
    for doc_data in matched_docs:
        doc.add_paragraph(f"- {doc_data['title']} (유사도: {int(doc_data['similarity']*100)}%)")
        doc.add_paragraph(f"  ⤷ 출처: {doc_data.get('source', '출처 미확인')}")
        doc.add_paragraph(f"  ⤷ 설명: {doc_data.get('full_context', doc_data['snippet'])}")
    doc.add_heading("AI 응답", level=2)
    doc.add_paragraph(llm_answer)
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

def run_rag_from_input(eGFR, creatinine, albumin, proteinuria):
    user_question = f"eGFR: {eGFR}, Creatinine: {creatinine}, Albumin: {albumin}, 단백뇨: {proteinuria} 이 수치로 어느 질병이 의심됩니까?"
    st.subheader("🔍 RAG 기반 결과 분석")
    st.write(f"❓ 자동 생성 질의: {user_question}")

    model = SentenceTransformer('sentence-transformers/distiluse-base-multilingual-cased-v2')
    question_embedding = model.encode([user_question])
    index = faiss.read_index("nephro_faiss.index")

    with open("nephro_chunks.json", "r", encoding="utf-8") as f:
        chunks = json.load(f)

    D, I = index.search(question_embedding, k=2)
    matched_docs = []
    for idx, score in zip(I[0], D[0]):
        chunk = chunks[idx]
        matched_docs.append({
            "title": f"관련 문서 #{idx+1}",
            "similarity": round(1.0 / (1.0 + score), 2),
            "snippet": chunk.get("text", ""),
            "full_context": chunk.get("text", ""),
            "source": chunk.get("source", "출처 없음")
        })

    st.markdown("### 🔗 검사된 문서 및 유사도")
    for doc in matched_docs:
        st.write(f"📄 **{doc['title']}** — 유사도: {int(doc['similarity'] * 100)}%")
        st.caption(f"➤️ 출처: {doc['source']}")
        st.caption(f"➤️ {doc['snippet']}")

    llm_answer = "(예시 응답) 입력된 수치는 CKD, Nephrotic Syndrome 등과 관련이 있을 수 있습니다. 정확한 진단은 의료진 상담이 필요합니다."
    st.markdown("### 🤖 AI 응답 예시")
    st.success(llm_answer)

    report_buffer = generate_report(user_question, matched_docs, llm_answer)
    report_bytes = report_buffer.read()
    st.download_button("📥 리포트 다운로드", data=report_bytes, file_name="질병예측_리포트.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

if submitted:
    errors = validate_input(eGFR, creatinine, albumin, proteinuria)
    if errors:
        for error in errors:
            st.warning(error)
        st.stop()

    st.subheader("🩺 예비 질병 결과")
    result = []
    if eGFR < 60:
        result.append("🔸 CKD 가능성 있음 (eGFR < 60)")
    if creatinine > 1.3:
        result.append("🔸 신기능 저하 가능성 (Creatinine > 1.3)")
    if albumin < 3.5:
        result.append("🔸 저알부민혈증 가능성")
    if proteinuria > 150:
        result.append("🔸 단백뇨 의심 (정상 기준 초과)")
    if result:
        for r in result:
            st.write(r)
    else:
        st.success("정상 범위로 판단됩니다.")

    run_rag_from_input(eGFR, creatinine, albumin, proteinuria)
